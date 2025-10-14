import uuid
import os
from rest_framework import viewsets, permissions, status
from rest_framework.response import Response
from rest_framework.views import APIView
from .models import Project, RFPDocument, QuestionAnswer
from .serializers import (
    ProjectSerializer,
    RFPDocumentSerializer,
    PromptSerializer,
    ProjectRAGSerializer,
)
from .utils.rag_service import RAGService
from .utils.openai_setup import chat_with_gpt
from langchain_community.document_loaders import Docx2txtLoader, PyPDFLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from django.conf import settings
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
from django.http import HttpResponse



class ProjectViewSet(viewsets.ModelViewSet):
    """
    A ViewSet for viewing and editing Project instances.
    """

    queryset = Project.objects.all()
    serializer_class = ProjectSerializer

    permission_classes = [permissions.IsAuthenticated]

    def perform_create(self, serializer):
        if "manager" not in serializer.validated_data:
            serializer.save(manager=self.request.user)
        else:
            serializer.save()


class RFPDocumentViewSet(viewsets.ModelViewSet):
    queryset = RFPDocument.objects.all()
    serializer_class = RFPDocumentSerializer
    permission_classes = [permissions.IsAuthenticated]

    def perform_create(self, serializer):
        uploaded_file = self.request.data.get("document_file")
        original_filename = uploaded_file.name

        base_name, file_extension = os.path.splitext(original_filename)

        serializer.save(
            uploaded_by=self.request.user,
            file_id=str(uuid.uuid4()),
            filename=base_name,
            file_type=file_extension.lstrip(".").lower(),
        )

    def perform_update(self, serializer):
        uploaded_file = self.request.data.get("document_file")

        update_kwargs = {}

        if uploaded_file:
            original_filename = uploaded_file.name
            base_name, file_extension = os.path.splitext(original_filename)

            update_kwargs['filename'] = base_name
            update_kwargs['file_type'] = file_extension.lstrip(".").lower()

        # Save the instance with any collected updates
        serializer.save(**update_kwargs)


class QueryRAGView(APIView):
    """
    POST endpoint to accept a prompt, query the RAG service, and return the answer.
    """

    permission_classes = [permissions.IsAuthenticated]

    def post(self, request):
        serializer = PromptSerializer(data=request.data)

        if serializer.is_valid():
            user_query = serializer.validated_data["prompt"]
            project_id = serializer.validated_data["project_id"]

            try:
                # Check if project exists
                try:
                    project = Project.objects.get(id=project_id)
                except Project.DoesNotExist:
                    return Response(
                        {"error": f"Project with ID {project_id} not found."},
                        status=status.HTTP_404_NOT_FOUND,
                    )

                # Get the RAG chain (initialized only the first time)
                rag_chain = RAGService.get_rag_chain()

                print(f"Invoking RAG chain for query: '{user_query}'")
                answer = rag_chain.invoke(user_query)

                # Save Q&A (use the project we already fetched above)
                QuestionAnswer.objects.create(
                    project=project,  # Use the project from above
                    question=user_query,
                    answer=answer
                )

                return Response(
                    {"query": user_query, "answer": answer}, 
                    status=status.HTTP_200_OK
                )

            except Exception as e:
                print(f"RAG Chain Error: {e}")
                return Response(
                    {
                        "error": "An error occurred while querying the knowledge base.",
                        "details": str(e),
                    },
                    status=status.HTTP_500_INTERNAL_SERVER_ERROR,
                )

        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

class InsertRAGView(APIView):
    """
    POST endpoint to take a project ID, retrieve associated documents,
    process, and insert chunks into Pinecone.
    """

    permission_classes = [permissions.IsAuthenticated]

    def post(self, request):
        serializer = ProjectRAGSerializer(data=request.data)

        if not serializer.is_valid():
            return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

        project_id = serializer.validated_data["project_id"]

        try:
            project = Project.objects.get(pk=project_id)

            if not project:
                return Response(
                    {"error": f"Project with ID {project_id} not found."},
                    status=status.HTTP_404_NOT_FOUND,
                )

            rfp_documents = project.documents.all()
            if not rfp_documents:
                return Response(
                    {
                        "message": f"No RFPDocuments found for Project ID {project_id}. Nothing to index."
                    },
                    status=status.HTTP_200_OK,
                )

            text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=1000, chunk_overlap=200
            )

            all_chunked_docs = []

            for doc in rfp_documents:
                # This ensures the path is correct regardless of OS
                file_path = Path(settings.MEDIA_ROOT) / doc.document_file.name

                if not file_path.exists():
                    print(
                        f"File not found for document {doc.filename}: {file_path}")
                    # Optionally skip this document or return an error
                    continue

                print(
                    f"Processing document: {doc.filename} at {file_path} with {doc.file_type}")

                if (doc.file_type == 'pdf'):
                    loader = PyPDFLoader(
                        str(file_path)
                    )
                elif (doc.file_type == 'docx'):
                    loader = Docx2txtLoader(
                        str(file_path)
                    )
                else:
                    print(
                        f"Unsupported file_type '{doc.file_type}' for document. Skipping.")
                    continue
                documents = loader.load()

                chunked_docs = text_splitter.split_documents(documents)
                print(f"Doc Split into {len(chunked_docs)} chunks.")

                all_chunked_docs.extend(chunked_docs)

            if not all_chunked_docs:
                return Response(
                    {
                        "message": f"No processable chunks found for Project ID {project_id}. Nothing to insert."
                    },
                    status=status.HTTP_200_OK,
                )

            # 4. Insert Data into Pinecone using RAGService
            result = RAGService.insert_documents(all_chunked_docs)

            return Response(
                {
                    "message": f"Successfully inserted {result.get('inserted_count', 0)} chunks for project {project_id}.",
                    "project_id": project_id,
                    "documents_processed": len(rfp_documents),
                },
                status=status.HTTP_200_OK,
            )

        except Project.DoesNotExist:
            return Response(
                {"error": f"Project with ID {project_id} not found in the database."},
                status=status.HTTP_404_NOT_FOUND,
            )
        except Exception as e:
            # Catch file access, database connection, or LLM errors
            print(f"RAG Insertion Error: {e}")
            return Response(
                {
                    "error": "An error occurred during data insertion.",
                    "details": str(e),
                },
                status=status.HTTP_500_INTERNAL_SERVER_ERROR,
            )


class OpenAIChat(APIView):
    def post(self, request):
        messages = request.data.get("messages", [])

        if not isinstance(messages, list) or not messages:
            return Response(
                {"error": "messages must be a non-empty list"},
                status=status.HTTP_400_BAD_REQUEST
            )

        try:
            answer = chat_with_gpt(messages)
            return Response({"response": answer}, status=status.HTTP_200_OK)
        except Exception as e:
            return Response(
                {"error": str(e)},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )


class GenerateRFPView(APIView):
    """
    POST endpoint to generate an AI-powered RFP document for a given project.
    Accepts project_id and returns a downloadable .docx file.
    """

    permission_classes = [permissions.IsAuthenticated]

    def get(self, request):
        """Handle browser's automatic GET request after download"""
        return Response(
            {"message": "RFP generation endpoint. Use POST with project_id to generate documents."},
            status=status.HTTP_200_OK
        )

    def post(self, request):
        serializer = ProjectRAGSerializer(data=request.data)

        if not serializer.is_valid():
            return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

        project_id = serializer.validated_data["project_id"]

        try:
            # Fetch the project
            project = Project.objects.get(pk=project_id)

            # Fetch unique Q&A pairs 
            qa_pairs = []
            seen_questions = set()

            all_qas = QuestionAnswer.objects.filter(
                project=project
            ).order_by('-created_at')

            for qa in all_qas:
                question_normalized = qa.question.strip().lower()
                
                # Skip if question already seen
                if question_normalized in seen_questions:
                    continue
                
                # Skip if answer is unhelpful
                if qa.answer.strip() == "I don't have enough information to answer that.":
                    print(f"Skipping Q&A with no answer: {qa.question[:50]}...")
                    continue
                
                # Only add if it passes both checks
                qa_pairs.append(qa)
                seen_questions.add(question_normalized)
                

            if not qa_pairs:
                return Response(
                    {
                        "error": "No Q&A history found for this project. "
                                "Please query the RAG system first to gather information."
                    },
                    status=status.HTTP_400_BAD_REQUEST,
                )

            # Build context from Q&A pairs
            context = self._build_context_from_qa(qa_pairs)

            # Generate comprehensive RFP content using GPT
            print(f"Generating AI-powered RFP content for project: {project.name}")
            rfp_content = self._generate_rfp_content_with_gpt(project, context)

            # Create formatted Word document
            doc = self._create_word_document(project, rfp_content, qa_pairs)

            # Save document to BytesIO buffer
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            # --- Save to server (media/generated_rfps/YYYY/MM/DD/) ---
            # Generate filename
            filename = f'{project.name.replace(" ", "_")}_RFP_Response.docx'
            
            # Create directory structure: generated_rfps/2025/10/11/
            current_date = datetime.now()
            generated_rfps_dir = Path(settings.MEDIA_ROOT) / 'generated_rfps' / \
                                current_date.strftime('%Y') / \
                                current_date.strftime('%m') / \
                                current_date.strftime('%d')
            
            # Create all necessary directories
            generated_rfps_dir.mkdir(parents=True, exist_ok=True)

            # Full path for the file
            file_path = generated_rfps_dir / filename

            # Save the document to the server
            with open(file_path, 'wb') as f:
                f.write(buffer.getvalue())

            print(f"✅ RFP saved to: {file_path}")

            # Reset buffer for HTTP response
            buffer.seek(0)

            # Create HTTP response for download
            response = HttpResponse(
                buffer.getvalue(),
                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
            
            response['Content-Disposition'] = f'attachment; filename="{filename}"'

            return response

        except Project.DoesNotExist:
            return Response(
                {"error": f"Project with ID {project_id} not found."},
                status=status.HTTP_404_NOT_FOUND,
            )
        except Exception as e:
            print(f"RFP Generation Error: {e}")
            import traceback
            traceback.print_exc()
            return Response(
                {
                    "error": "An error occurred while generating the RFP document.",
                    "details": str(e),
                },
                status=status.HTTP_500_INTERNAL_SERVER_ERROR,
            )

    def _build_context_from_qa(self, qa_pairs):
        """
        Builds a context string from all Q&A pairs for GPT to analyze.
        """
        context_parts = []
        for idx, qa in enumerate(qa_pairs, 1):
            context_parts.append(f"Q{idx}: {qa.question}")
            context_parts.append(f"A{idx}: {qa.answer}")
            context_parts.append("---")
        
        return "\n".join(context_parts)

    def _generate_rfp_content_with_gpt(self, project, context):
        """
        Uses GPT to generate a comprehensive, professional RFP response document
        based on project details and Q&A context.
        """
        system_prompt = """You are an expert RFP (Request for Proposal) response writer. 
            Create a professional, comprehensive RFP response document with these sections:

            1. **Company Introduction** - Overview of the company and capabilities
            2. **Product** - Detailed product/solution description
            3. **Support** - Support services and maintenance offered
            4. **Contract Policies** - Terms, conditions, and policies
            5. **Pricing** - Cost breakdown and payment terms
            6. **Reference** - Past projects, case studies, or client references

            Guidelines:
            - Use ONLY information from the provided Q&A context
            - If a section has no relevant information in the Q&A, skip that section entirely
            - Do NOT make up or assume any information
            - Write in a professional, confident tone
            - Use clear section headers with ##
            - Keep content factual and specific"""

        user_prompt = f"""Generate an RFP response document for this project:
                **Project Details:**
                - Name: {project.name}
                - Type: {project.type}
                - Description: {project.description}
                - Due Date: {project.due_date.strftime('%B %d, %Y')}
                - Value: ${project.value:,}

                **Q&A Context:**
                {context}

                Create the RFP response using ONLY the information above. Skip any section that has no relevant data in the Q&A context.
            """

        messages = [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt}
        ]

        print("Calling GPT to generate RFP content...")
        rfp_content = chat_with_gpt(messages)
        print("GPT generation complete.")
        
        return rfp_content

    def _create_word_document(self, project, ai_generated_content, qa_pairs):
        """
        Creates a beautifully formatted Word document with AI-generated content
        plus the original Q&A as an appendix.
        """
        doc = Document()

        # ============== COVER PAGE ==============
        # Title
        title = doc.add_heading(f'Request for Proposal Response', level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        subtitle = doc.add_heading(project.name, level=1)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()  # Spacing

        # Project Metadata Table
        details_table = doc.add_table(rows=6, cols=2)
        details_table.style = 'Light Grid Accent 1'
        
        details_data = [
            ('Project Type', project.type),
            ('Due Date', project.due_date.strftime('%B %d, %Y')),
            ('Project Value', f'${project.value:,}'),
            ('Current Stage', project.get_stage_display()),
            ('Project Manager', project.manager.email),
            ('Generated On', datetime.now().strftime('%B %d, %Y at %I:%M %p')),
        ]

        for i, (label, value) in enumerate(details_data):
            row = details_table.rows[i]
            row.cells[0].text = label
            row.cells[1].text = str(value)
            row.cells[0].paragraphs[0].runs[0].font.bold = True

        doc.add_page_break()

        # ============== AI-GENERATED CONTENT ==============
        # Parse and format the AI-generated content
        lines = ai_generated_content.split('\n')
        
        for line in lines:
            line = line.strip()
            
            if not line:
                continue
            
            # Check for markdown headers
            if line.startswith('## '):
                # Main section header
                heading_text = line.replace('## ', '').strip()
                doc.add_heading(heading_text, level=1)
            elif line.startswith('### '):
                # Sub-section header
                heading_text = line.replace('### ', '').strip()
                doc.add_heading(heading_text, level=2)
            elif line.startswith('**') and line.endswith('**'):
                # Bold text (likely a sub-heading)
                bold_text = line.strip('**')
                p = doc.add_paragraph()
                p.add_run(bold_text).bold = True
            elif line.startswith('- ') or line.startswith('* '):
                # Bullet point
                bullet_text = line[2:].strip()
                doc.add_paragraph(bullet_text, style='List Bullet')
            elif line.startswith(tuple(f'{i}.' for i in range(1, 10))):
                # Numbered list
                doc.add_paragraph(line, style='List Number')
            else:
                # Regular paragraph
                doc.add_paragraph(line)

        doc.add_page_break()
        doc.add_heading('Appendix A: Question & Answer Reference', level=1)
        
        appendix_intro = doc.add_paragraph(
            'The following questions and answers were gathered during the discovery phase '
            'and informed the content of this proposal.'
        )
        appendix_intro.italic = True
        
        doc.add_paragraph()

        for idx, qa in enumerate(qa_pairs, 1):
            # Question
            q_paragraph = doc.add_paragraph()
            q_run = q_paragraph.add_run(f'Q{idx}: {qa.question}')
            q_run.font.bold = True
            q_run.font.size = Pt(11)
            
            # Answer
            a_paragraph = doc.add_paragraph()
            a_run = a_paragraph.add_run(f'A: {qa.answer}')
            a_run.font.size = Pt(10)
            
            doc.add_paragraph()  # Spacing

        doc.add_page_break()
        footer_para = doc.add_paragraph()
        footer_run = footer_para.add_run(
            f'\nThis document was generated using AI analysis of project requirements.\n'
            f'Generated on {datetime.now().strftime("%B %d, %Y at %I:%M %p")}'
        )
        footer_run.font.size = Pt(9)
        footer_run.font.color.rgb = RGBColor(128, 128, 128)
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        return doc

